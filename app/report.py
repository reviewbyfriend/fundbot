from io import BytesIO
from decimal import Decimal
from pathlib import Path
import re
import shutil
import tempfile

from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from .services import money
from .timezone import format_th, format_now_th, now_bangkok

TEMPLATE_PATH = Path(__file__).parent / "static" / "report_template.xlsx"

THAI_MONTHS_FULL = {
    1: "มกราคม", 2: "กุมภาพันธ์", 3: "มีนาคม", 4: "เมษายน", 5: "พฤษภาคม", 6: "มิถุนายน",
    7: "กรกฎาคม", 8: "สิงหาคม", 9: "กันยายน", 10: "ตุลาคม", 11: "พฤศจิกายน", 12: "ธันวาคม",
}
THAI_MONTHS_SHORT = {
    1: "ม.ค.", 2: "ก.พ.", 3: "มี.ค.", 4: "เม.ย.", 5: "พ.ค.", 6: "มิ.ย.",
    7: "ก.ค.", 8: "ส.ค.", 9: "ก.ย.", 10: "ต.ค.", 11: "พ.ย.", 12: "ธ.ค.",
}


def _status_th(p):
    if p.status == "paid":
        if getattr(p, "payment_type", "") == "cash":
            return "ชำระแล้ว (เงินสด)"
        return "ชำระแล้ว (โอน)"
    if p.status == "pending":
        return "รอตรวจสอบ"
    return "ยังไม่ได้ชำระ"


def _expense_date(e):
    return format_th(getattr(e, "expense_date", None) or getattr(e, "created_at", None), "%d/%m/%Y", default="")


def _d(x):
    try:
        return Decimal(x or 0)
    except Exception:
        return Decimal("0")


def report_summary(round_, payments, expenses=None):
    expenses = expenses or []
    carry = _d(getattr(round_, "carry_over", 0))
    total_due = sum([_d(p.due_amount) for p in payments], Decimal("0"))
    # แบบรายงานสำนักงานใช้ยอดรับรายเดือนตามรายชื่อ (จำนวนค่อนข้างคงที่)
    total_income = carry + total_due
    total_paid = sum([_d(p.due_amount) for p in payments if p.status == "paid"], Decimal("0"))
    total_expense = sum([_d(e.amount) for e in expenses], Decimal("0"))
    balance = total_income - total_expense
    return {
        "carry_over": carry,
        "monthly_income": total_due,
        "total_income": total_income,
        "paid_actual": total_paid,
        "total_expense": total_expense,
        "balance": balance,
        "expense_count": len(expenses),
        "paid_count": len([p for p in payments if p.status == "paid"]),
        "pending_count": len([p for p in payments if p.status == "pending"]),
        "unpaid_count": len([p for p in payments if p.status not in ["paid", "pending"]]),
    }


def _month_from_round(round_):
    title = (getattr(round_, "title", "") or "").strip()
    # 2026-06 / 2569-06
    m = re.search(r"(20\d{2}|25\d{2})[-/](\d{1,2})", title)
    if m:
        return int(m.group(2))
    # Thai month names in title
    for i, name in THAI_MONTHS_FULL.items():
        if name in title:
            return i
    for i, name in THAI_MONTHS_SHORT.items():
        if name.replace('.', '') in title.replace('.', ''):
            return i
    return now_bangkok().month


def _select_template_sheet(wb, month: int):
    full = THAI_MONTHS_FULL[month]
    short = THAI_MONTHS_SHORT[month].replace('.', '')
    # prefer exact current form if exists
    for ws in wb.worksheets:
        t = ws.title.replace('.', '')
        if full in (ws["A2"].value or "") or short in t:
            return ws
    return wb["Sheet1"] if "Sheet1" in wb.sheetnames else wb.active


def _date_text(dt):
    return format_th(dt, "%d %b %y", default="")


def _copy_sheet_only(wb, ws):
    # ลบชีตอื่นทั้งหมด เหลือฟอร์มเดียวสำหรับพิมพ์
    for other in list(wb.worksheets):
        if other != ws:
            wb.remove(other)
    ws.title = "รายงาน"
    return ws


def _report_rows(ws):
    # ตรวจจากฟอร์ม: รายจ่ายอยู่ C/D/E เริ่มแถว 4 ถึงก่อนแถวรวมรายจ่าย
    total_row = None
    for row in range(4, ws.max_row + 1):
        if str(ws.cell(row, 4).value or "").strip() == "รวมรายจ่าย":
            total_row = row
            break
    if not total_row:
        total_row = 12
    return 4, total_row - 1, total_row


def _summary_rows(ws):
    income_row = expense_row = balance_row = None
    for row in range(1, ws.max_row + 1):
        label = str(ws.cell(row, 4).value or "").strip()
        if label == "รายรับ" and income_row is None:
            income_row = row
        elif label == "รายจ่าย" and expense_row is None:
            expense_row = row
        elif "เงินคงเหลือ" in label:
            balance_row = row
    return income_row, expense_row, balance_row


def _clear_expense_area(ws, start_row, end_row):
    for row in range(start_row, end_row + 1):
        ws.cell(row, 3).value = None
        ws.cell(row, 4).value = None
        ws.cell(row, 5).value = None


def _write_template_report(wb, ws, round_, payments, expenses=None):
    expenses = sorted(expenses or [], key=lambda e: (getattr(e, "expense_date", None) or getattr(e, "created_at", None), getattr(e, "id", 0)))
    month = _month_from_round(round_)
    year_be = now_bangkok().year + 543
    ws["A2"] = f"ประจำเดือน {THAI_MONTHS_FULL[month]} {year_be}"

    # รายรับ: คงรูปฟอร์มเดิม เปลี่ยนเฉพาะตัวเลขด้านใน
    start_row, end_row, total_row = _report_rows(ws)
    carry = _d(getattr(round_, "carry_over", 0))
    ws["B4"] = float(carry)
    # ใส่ยอดสมาชิกตามชื่อในฟอร์มเดิม ถ้าหาชื่อเจอ
    pay_by_name = {getattr(p.member, "name", "").replace(" ", ""): _d(p.due_amount) for p in payments}
    used = set()
    for row in range(5, total_row):
        name = str(ws.cell(row, 1).value or "").replace(" ", "")
        if not name:
            continue
        match_amount = None
        for key, amount in pay_by_name.items():
            if key and (key in name or name in key):
                match_amount = amount
                used.add(key)
                break
        if match_amount is not None:
            ws.cell(row, 2).value = float(match_amount)

    # รายจ่าย: เขียนทับเฉพาะช่องวันที่/รายการ/ยอด ในพื้นที่เดิม
    _clear_expense_area(ws, start_row, end_row)
    max_items = end_row - start_row + 1
    for idx, e in enumerate(expenses[:max_items]):
        row = start_row + idx
        ws.cell(row, 3).value = _date_text(getattr(e, "expense_date", None) or getattr(e, "created_at", None))
        ws.cell(row, 4).value = getattr(e, "title", "")
        ws.cell(row, 5).value = float(_d(getattr(e, "amount", 0)))
    if len(expenses) > max_items:
        # ไม่เปลี่ยน layout: รวมรายการเกินในแถวสุดท้ายเป็น "อื่น ๆ"
        extra = expenses[max_items-1:]
        row = end_row
        ws.cell(row, 3).value = ""
        ws.cell(row, 4).value = f"อื่น ๆ รวม {len(extra)} รายการ"
        ws.cell(row, 5).value = float(sum([_d(e.amount) for e in extra], Decimal("0")))

    # คงสูตรเดิมของฟอร์ม แต่ปรับช่วงรวมให้ตรงพื้นที่เดิม
    ws.cell(total_row, 2).value = f"=SUM(B4:B{total_row-1})"
    ws.cell(total_row, 5).value = f"=SUM(E{start_row}:E{end_row})"

    summary = report_summary(round_, payments, expenses)
    income_row, expense_row, balance_row = _summary_rows(ws)
    if income_row:
        ws.cell(income_row, 5).value = float(summary["total_income"])
    if expense_row:
        ws.cell(expense_row, 5).value = float(summary["total_expense"])
    if balance_row:
        ws.cell(balance_row, 5).value = float(summary["balance"])

    # วันที่ลงท้าย: เปลี่ยนเฉพาะช่องวันที่ถ้ามี
    today = format_now_th("%d %b %y")
    for row in range(1, ws.max_row + 1):
        for col in [1, 4]:
            v = str(ws.cell(row, col).value or "")
            if re.search(r"\d{1,2}\s*(ม|ก|พ|เม|มิ|ส|ต|ธ|ก\.พ|มี\.ค|เม\.ย|มิ\.ย|ส\.ค|ก\.ย|ต\.ค|พ\.ย|ธ\.ค)", v):
                # แก้เฉพาะช่องที่น่าจะเป็นวันที่ลายเซ็นท้ายเอกสาร
                if row >= 20:
                    ws.cell(row, col).value = today
    return ws


def make_excel(round_, payments, expenses=None):
    expenses = expenses or []
    if TEMPLATE_PATH.exists():
        # ใช้แม่แบบ Excel ของสำนักงาน ห้ามเปลี่ยนฟอร์ม เปลี่ยนเฉพาะตัวเลข/ข้อมูลในช่องเท่านั้น
        wb = load_workbook(TEMPLATE_PATH)
        ws = _select_template_sheet(wb, _month_from_round(round_))
        ws = _copy_sheet_only(wb, ws)
        _write_template_report(wb, ws, round_, payments, expenses)
        out = BytesIO()
        wb.save(out)
        return out.getvalue()

    # fallback เท่านั้น กรณี template หาย
    wb = Workbook()
    ws = wb.active
    ws.title = "สรุปรายเดือน"
    blue = "1F6FEB"
    red = "D93025"
    border = Border(bottom=Side(style="thin", color="E5E7EB"))
    ws["A1"] = "รายงานเงินกองสำนักงาน"
    ws["A1"].font = Font(size=18, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=blue)
    ws.merge_cells("A1:F1")
    ws["A2"] = f"รอบเดือน: {round_.title}"
    ws["A3"] = f"สร้างรายงาน: {format_now_th('%d/%m/%Y %H:%M')}"
    s = report_summary(round_, payments, expenses)
    for row, (label, val) in enumerate([("ยอดยกมา", s["carry_over"]), ("รวมรายรับ", s["total_income"]), ("รวมรายจ่าย", s["total_expense"]), ("เงินคงเหลือ", s["balance"])], 5):
        ws.cell(row, 1, label).font = Font(bold=True)
        ws.cell(row, 2, float(val)).number_format = '#,##0.00'
    out = BytesIO(); wb.save(out); return out.getvalue()


def make_word(round_, payments, expenses=None):
    expenses = expenses or []
    doc = Document()
    doc.add_heading("รายงานเงินกองสำนักงาน", 0)
    doc.add_paragraph(f"รอบเดือน: {round_.title}")
    doc.add_paragraph(f"สร้างรายงาน: {format_now_th('%d/%m/%Y %H:%M')}")
    s = report_summary(round_, payments, expenses)
    doc.add_paragraph(f"รายรับรวม: {money(s['total_income'])} บาท")
    doc.add_paragraph(f"รายจ่ายรวม: {money(s['total_expense'])} บาท")
    doc.add_paragraph(f"คงเหลือ: {money(s['balance'])} บาท")
    doc.add_paragraph("หมายเหตุ: ไฟล์ Excel เป็นไฟล์หลักตามแบบฟอร์มสำนักงาน ใช้สำหรับพิมพ์เอกสาร")
    doc.add_heading("รายจ่าย", level=1)
    et = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["วันที่", "รายการ", "ยอด", "หมายเหตุ"]):
        et.rows[0].cells[i].text = h
    for e in expenses:
        row = et.add_row().cells
        row[0].text = _expense_date(e)
        row[1].text = e.title
        row[2].text = money(e.amount)
        row[3].text = getattr(e, "note", "") or ""
    out = BytesIO(); doc.save(out); return out.getvalue()


def make_pdf(round_, payments, expenses=None):
    expenses = expenses or []
    out = BytesIO()
    c = canvas.Canvas(out, pagesize=A4)
    w, h = A4
    y = h - 50
    s = report_summary(round_, payments, expenses)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "FundBot Monthly Summary")
    y -= 24
    c.setFont("Helvetica", 11)
    c.drawString(40, y, f"Round: {round_.title}")
    y -= 18
    c.drawString(40, y, f"Income: {money(s['total_income'])}  Expense: {money(s['total_expense'])}  Balance: {money(s['balance'])}")
    y -= 28
    c.setFont("Helvetica-Bold", 12)
    c.drawString(40, y, "Expenses")
    y -= 18
    c.setFont("Helvetica", 10)
    for e in expenses:
        if y < 80:
            c.showPage(); y = h - 50; c.setFont("Helvetica", 10)
        c.drawString(40, y, _expense_date(e))
        c.drawString(110, y, str(e.title)[:36])
        c.drawString(360, y, money(e.amount))
        y -= 15
    c.save()
    return out.getvalue()
